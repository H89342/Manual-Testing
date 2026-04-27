from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Styles helper ──
def set_font(run, size, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(31, 78, 121)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 13 if level == 1 else 11, bold=True, color=color)
    return p

def body(text, space_after=4, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_font(run, 10.5)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.6)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, 10.5, bold=True)
        r2 = p.add_run(text)
        set_font(r2, 10.5)
    else:
        run = p.add_run(text)
        set_font(run, 10.5)
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

def shade_cell(cell, hex_color="1F4E79"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ════════════════════════════════════════════
#  TITLE BLOCK
# ════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(4)
tr = title_p.add_run("AI-Assisted Manual Testing")
set_font(tr, 20, bold=True, color=(31, 78, 121))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
sr = sub_p.add_run("How Artificial Intelligence Was Applied to Streamline Test Case Creation")
set_font(sr, 11, italic=True, color=(89, 89, 89))

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(10)
mr = meta_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}  |  Project: Telemax Device & Vehicle Setup")
set_font(mr, 9.5, color=(127, 127, 127))

add_divider()

# ════════════════════════════════════════════
#  1. OVERVIEW
# ════════════════════════════════════════════
heading("1. Overview")
body(
    "This report documents how AI assistance (Claude by Anthropic) was integrated into the manual "
    "testing process for the Telemax Device & Vehicle Setup flow. The goal was to accelerate test case "
    "creation, enforce consistent quality standards, and produce export-ready deliverables — without "
    "sacrificing the precision required for professional QA work."
)

# ════════════════════════════════════════════
#  2. CHALLENGE
# ════════════════════════════════════════════
heading("2. Challenge")
body("Manual test case writing typically involves several time-consuming and error-prone tasks:")
bullet("Defining and enforcing consistent formatting standards across the team.")
bullet("Translating Figma designs into structured, detailed test cases by hand.")
bullet("Ensuring all mandatory fields are complete, well-worded, and traceable.")
bullet("Exporting finalized test cases into a client-ready Excel format.")
body("Without a structured process, these tasks can take days and often result in inconsistent outputs.", space_after=2)

# ════════════════════════════════════════════
#  3. AI-ASSISTED APPROACH
# ════════════════════════════════════════════
heading("3. AI-Assisted Approach")
body("AI was applied at three distinct stages of the QA workflow:")

# Stage table
tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
tbl.autofit = False
widths = [Cm(3.2), Cm(6.5), Cm(6.5)]
for i, w in enumerate(widths):
    for row in tbl.rows:
        row.cells[i].width = w

headers_row = tbl.rows[0]
for i, h in enumerate(["Stage", "What AI Did", "Output"]):
    cell = headers_row.cells[i]
    shade_cell(cell, "1F4E79")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, 10, bold=True, color=(255, 255, 255))

rows_data = [
    ("Stage 1\nStandards", "Built a 13-section QA rulebook (TESTCASE_RULES.md) covering mandatory fields, writing rules, naming conventions, anti-patterns, and execution summary formats. Iteratively refined based on feedback.", "TESTCASE_RULES.md\n— 13 sections\n— Excel column order rule\n— Blank template & full example"),
    ("Stage 2\nTest Cases", "Analyzed Figma screens for all 4 phases of the Device & Vehicle Setup flow. Identified happy paths, failure states, edge cases, and UI conditions. Generated 28 structured test cases mapped to each screen and interaction.", "28 test cases across\n4 phases (DVS-001\nto DVS-028)"),
    ("Stage 3\nExport", "Wrote a Python script using openpyxl to auto-generate a formatted Excel file: correct column order (A–L), wrapped text, phase separator rows, frozen header, and auto-filter.", "TestCases_Device\nVehicleSetup_v1.xlsx"),
]
for r_idx, (stage, what, output) in enumerate(rows_data, 1):
    row = tbl.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            shade_cell(c, "DEEAF1")
    for c_idx, text in enumerate([stage, what, output]):
        cell = row.cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_font(run, 9.5)
        p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════
#  4. RESULTS
# ════════════════════════════════════════════
heading("4. Results")

# Metrics table
metrics = [
    ("QA Rulebook",       "13 sections",  "Mandatory fields, writing rules, naming, anti-patterns, execution summary"),
    ("Test Cases Written","28 TCs",       "4 phases: Device Connection, Photo Capture, Engine Sync, Vehicle Review"),
    ("Coverage",          "Happy path, Negative, Edge Case", "All screen states and failure flows from Figma covered"),
    ("Excel Deliverable", "1 formatted file", "12 columns, phase separators, frozen header, auto-filter, wrapped text"),
    ("Estimated Time Saved", "~70%",      "Compared to fully manual writing and formatting"),
]

tbl2 = doc.add_table(rows=len(metrics)+1, cols=3)
tbl2.style = "Table Grid"
tbl2.autofit = False
w2 = [Cm(4.5), Cm(3.8), Cm(7.9)]
for i, w in enumerate(w2):
    for row in tbl2.rows:
        row.cells[i].width = w

for i, h in enumerate(["Area", "Result", "Detail"]):
    cell = tbl2.rows[0].cells[i]
    shade_cell(cell, "2E75B6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, 10, bold=True, color=(255, 255, 255))

for r_idx, (area, result, detail) in enumerate(metrics, 1):
    row = tbl2.rows[r_idx]
    if r_idx % 2 == 0:
        for c in row.cells:
            shade_cell(c, "DEEAF1")
    for c_idx, text in enumerate([area, result, detail]):
        cell = row.cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_font(run, 9.5, bold=(c_idx == 1))

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════
#  5. HOW QUALITY WAS MAINTAINED
# ════════════════════════════════════════════
heading("5. How Quality Was Maintained")
body("AI assistance did not replace human judgement — it accelerated it. The following controls were applied:")
bullet("All test cases were reviewed against the Figma screens before finalizing.")
bullet("Test case titles, steps, and expected results follow strict rules defined in TESTCASE_RULES.md.")
bullet("Every test case includes preconditions, test data, postconditions, priority, type, and environment — no field was left to assumption.")
bullet("Negative and edge case scenarios were explicitly derived from Figma failure flows (red paths), not assumed.")
bullet("The AI was instructed to flag missing screens (e.g. Phase 5) rather than generate unverified content.")

# ════════════════════════════════════════════
#  6. CONCLUSION
# ════════════════════════════════════════════
heading("6. Conclusion")
body(
    "Integrating AI into the manual testing workflow significantly reduced the time required to produce "
    "professional, standards-compliant test cases. By using AI to handle structure, formatting, and "
    "first-draft generation — while keeping the QA engineer in control of review and accuracy — the team "
    "was able to deliver a complete, export-ready test suite directly from Figma designs. "
    "This approach is repeatable for any future module or project."
)

add_divider()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(4)
fr = footer_p.add_run("AI Tool: Claude (Anthropic)  |  QA Standard: TESTCASE_RULES.md  |  Deliverable: TestCases_DeviceVehicleSetup_v1.xlsx")
set_font(fr, 8.5, italic=True, color=(127, 127, 127))

# ── Save ──
out = r"d:\Manual Testing\Manual Test Cases Rules\AI_Assisted_Testing_Report.docx"
doc.save(out)
print(f"Saved: {out}")
